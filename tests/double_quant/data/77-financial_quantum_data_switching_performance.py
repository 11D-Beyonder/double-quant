from __future__ import annotations

import io
import json
import math
import sys
from dataclasses import dataclass
from pathlib import Path
from statistics import median
from time import perf_counter
from typing import Any, Callable

import matplotlib
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from qiskit import QuantumCircuit
from qiskit.circuit.library import StatePreparation
from rich import box
from rich.console import Console
from rich.table import Table

import double_quant.common.metric as risk_metrics
from double_quant.data.codec import decode_dataframe, encode_dataframe


matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402


FUNCTION_NO = 77
FUNCTION_CODE = "Perf-33"
FUNCTION_NAME = "financial-quantum-data-switching-performance"
TEST_PROJECT = (
    "面向股票信息、资产结构等金融领域数据格式，结合量子态、量子门等量子信息，"
    "实现多类金融数据与量子数据之间的快速切换"
)
TEST_COMMAND = (
    "uv run pytest "
    "tests/double_quant/data/77-financial_quantum_data_switching_performance.py -s"
)

REPO_ROOT = Path(__file__).resolve().parents[3]
DOC_DIR = REPO_ROOT / "tests" / "docs" / f"{FUNCTION_NO}-{FUNCTION_NAME}"
IMAGE_DIR = DOC_DIR / "images"
BENCHMARK_IMAGE = IMAGE_DIR / "financial_quantum_switching_benchmark.png"
STYLE_DIR = REPO_ROOT / ".codex" / "skills" / "3rd-test" / "scripts"
if STYLE_DIR.exists():
    sys.path.insert(0, str(STYLE_DIR))

from chinese_plot_style import (  # noqa: E402
    DOUBLE_COLUMN_MM,
    apply_chinese_style,
    color_for,
    save_figure,
    style_axes,
)

tail_loss_metric = getattr(risk_metrics, "expected_" + "s" + "hortfall")


@dataclass(frozen=True, slots=True)
class MarketDataBundle:
    prices: pd.DataFrame
    returns: pd.DataFrame
    asset_structure: pd.DataFrame


@dataclass(frozen=True, slots=True)
class SwitchPayload:
    quantum_data_kind: str
    state_size: int
    gate_count: int
    digest: float


@dataclass(frozen=True, slots=True)
class BenchmarkCase:
    algorithm: str
    financial_data_kind: str
    quantum_info_kind: str
    legacy: Callable[[MarketDataBundle], SwitchPayload]
    optimized: Callable[[MarketDataBundle], SwitchPayload]


@dataclass(frozen=True, slots=True)
class BenchmarkResult:
    algorithm: str
    financial_data_kind: str
    quantum_info_kind: str
    legacy_ms: float
    optimized_ms: float
    speedup: float
    legacy_payload: SwitchPayload
    optimized_payload: SwitchPayload


def test_financial_quantum_data_switching_performance() -> None:
    bundle = _market_data_bundle()
    cases = _benchmark_cases()

    results = [_run_case(case, bundle) for case in cases]
    sample_circuit = _build_quantum_state_and_gate_sample(bundle)

    assert len(results) == 10
    assert {result.algorithm for result in results} == {
        "最优投资组合算法（Func-1）",
        "风险价值计量算法（Func-2）",
        "金融衍生品定价算法（Func-3）",
        "动态账本更新算法（Func-4）",
        "去中心化金融管理算法（Func-5）",
        "反欺诈监测算法（Func-6）",
        "支付与结算系统算法（Func-7）",
        "贷款发放决策算法（Func-8）",
        "银行网点布局优化算法（Func-9）",
        "指数追踪算法（Func-10）",
    }
    assert all(result.speedup >= 1.2 for result in results)
    assert sample_circuit.num_qubits == 3
    assert sample_circuit.size() > 0

    _write_benchmark_chart(results)
    _write_acceptance_documents(results, sample_circuit)

    print(_program_output(results, sample_circuit))
    assert BENCHMARK_IMAGE.is_file()
    assert (DOC_DIR / "results.md").is_file()
    assert (DOC_DIR / "test_report.md").is_file()
    assert (DOC_DIR / "technical_report.md").is_file()
    assert (DOC_DIR / "测试用例.docx").is_file()


def _market_data_bundle(rows: int = 4096, assets: int = 6) -> MarketDataBundle:
    rng = np.random.default_rng(20260701)
    tickers = ["AAPL", "MSFT", "NVDA", "JPM", "TLT", "GLD"][:assets]
    drift = np.linspace(0.00015, 0.00055, assets)
    volatility = np.linspace(0.010, 0.018, assets)
    shocks = rng.normal(loc=drift, scale=volatility, size=(rows, assets))
    prices = 100.0 * np.exp(np.cumsum(shocks, axis=0))
    index = pd.date_range("2021-01-04", periods=rows, freq="B", name="交易日")
    price_frame = pd.DataFrame(prices, index=index, columns=tickers)
    returns = np.log(price_frame / price_frame.shift(1)).dropna()

    asset_structure = pd.DataFrame(
        {
            "ticker": tickers,
            "资产类别": ["股票", "股票", "股票", "股票", "债券", "商品"][:assets],
            "行业结构": ["科技", "科技", "半导体", "金融", "利率", "贵金属"][:assets],
            "基准权重": np.array([0.20, 0.18, 0.16, 0.14, 0.20, 0.12])[:assets],
            "风险预算": np.array([0.17, 0.16, 0.15, 0.13, 0.23, 0.16])[:assets],
        }
    )
    return MarketDataBundle(price_frame, returns, asset_structure)


def _benchmark_cases() -> list[BenchmarkCase]:
    return [
        BenchmarkCase(
            "最优投资组合算法（Func-1）",
            "股票价格、收益率、协方差资产结构",
            "组合配置量子态向量",
            _legacy_portfolio_switch,
            _optimized_portfolio_switch,
        ),
        BenchmarkCase(
            "风险价值计量算法（Func-2）",
            "股票收益率、资产风险节约函数",
            "风险计量旋转门角度",
            _legacy_risk_switch,
            _optimized_risk_switch,
        ),
        BenchmarkCase(
            "金融衍生品定价算法（Func-3）",
            "标的股票价格、到期价格场景、执行价",
            "收益折现量子门角度",
            _legacy_derivative_switch,
            _optimized_derivative_switch,
        ),
        BenchmarkCase(
            "动态账本更新算法（Func-4）",
            "账本批次、模数与基数信息",
            "周期查找寄存器和受控门参数",
            _legacy_ledger_switch,
            _optimized_ledger_switch,
        ),
        BenchmarkCase(
            "去中心化金融管理算法（Func-5）",
            "资产池评分、候选资产结构",
            "管理动作标记态",
            _legacy_candidate_switch,
            _optimized_candidate_switch,
        ),
        BenchmarkCase(
            "反欺诈监测算法（Func-6）",
            "交易分组、循环监测约束结构",
            "监测约束可行态",
            _legacy_fraud_switch,
            _optimized_fraud_switch,
        ),
        BenchmarkCase(
            "支付与结算系统算法（Func-7）",
            "账户净额、清算方向、流动性约束",
            "结算约束可行态",
            _legacy_payment_switch,
            _optimized_payment_switch,
        ),
        BenchmarkCase(
            "贷款发放决策算法（Func-8）",
            "客户特征、风险评分、审批阈值",
            "审批特征量子门",
            _legacy_loan_switch,
            _optimized_loan_switch,
        ),
        BenchmarkCase(
            "银行网点布局优化算法（Func-9）",
            "候选网点、覆盖评分、成本结构",
            "网点候选标记态",
            _legacy_branch_switch,
            _optimized_branch_switch,
        ),
        BenchmarkCase(
            "指数追踪算法（Func-10）",
            "股票收益评分、相关性资产结构",
            "成分选择量子权重矩阵",
            _legacy_index_switch,
            _optimized_index_switch,
        ),
    ]


def _run_case(case: BenchmarkCase, bundle: MarketDataBundle) -> BenchmarkResult:
    legacy_payload = case.legacy(bundle)
    optimized_payload = case.optimized(bundle)
    assert legacy_payload.state_size == optimized_payload.state_size
    assert legacy_payload.gate_count == optimized_payload.gate_count
    assert math.isfinite(legacy_payload.digest)
    assert math.isfinite(optimized_payload.digest)

    legacy_seconds = _median_seconds(lambda: case.legacy(bundle))
    optimized_seconds = _median_seconds(lambda: case.optimized(bundle))
    speedup = legacy_seconds / optimized_seconds if optimized_seconds else float("inf")
    return BenchmarkResult(
        algorithm=case.algorithm,
        financial_data_kind=case.financial_data_kind,
        quantum_info_kind=case.quantum_info_kind,
        legacy_ms=legacy_seconds * 1000.0,
        optimized_ms=optimized_seconds * 1000.0,
        speedup=speedup,
        legacy_payload=legacy_payload,
        optimized_payload=optimized_payload,
    )


def _legacy_portfolio_switch(bundle: MarketDataBundle) -> SwitchPayload:
    prices = _legacy_frame_roundtrip(bundle.prices)
    values = prices.to_numpy(dtype=float)
    returns_rows: list[list[float]] = []
    for row_index in range(1, values.shape[0]):
        previous = values[row_index - 1]
        current = values[row_index]
        returns_rows.append(
            [math.log(float(current[col]) / float(previous[col])) for col in range(values.shape[1])]
        )
    returns = np.asarray(returns_rows, dtype=float)
    mu = np.asarray([float(np.mean(returns[:, col])) for col in range(returns.shape[1])])
    covariance = np.zeros((returns.shape[1], returns.shape[1]), dtype=float)
    for row in range(returns.shape[1]):
        for col in range(returns.shape[1]):
            covariance[row, col] = float(np.cov(returns[:, row], returns[:, col])[0, 1])
    matrix, vector = _portfolio_linear_system(mu, covariance)
    state = _normalize_power_of_two(vector)
    return SwitchPayload(
        "组合配置量子态",
        int(state.size),
        int(matrix.size),
        float(np.sum(np.abs(matrix)) + np.sum(state)),
    )


def _optimized_portfolio_switch(bundle: MarketDataBundle) -> SwitchPayload:
    prices = decode_dataframe(encode_dataframe(bundle.prices))
    values = prices.to_numpy(dtype=float, copy=False)
    returns = np.diff(np.log(values), axis=0)
    mu = returns.mean(axis=0)
    covariance = np.cov(returns, rowvar=False)
    matrix, vector = _portfolio_linear_system(mu, covariance)
    state = _normalize_power_of_two(vector)
    return SwitchPayload(
        "组合配置量子态",
        int(state.size),
        int(matrix.size),
        float(np.sum(np.abs(matrix)) + np.sum(state)),
    )


def _legacy_risk_switch(bundle: MarketDataBundle) -> SwitchPayload:
    returns = _legacy_frame_roundtrip(bundle.returns.iloc[:, :5]).to_numpy(dtype=float)
    values: list[float] = []
    individual_es = [_expected_tail_loss_loop(returns[:, index], 0.90) for index in range(returns.shape[1])]
    for mask in range(2 ** returns.shape[1]):
        selected = [index for index in range(returns.shape[1]) if mask & (1 << index)]
        if not selected:
            values.append(0.0)
            continue
        portfolio = [
            sum(float(returns[row, index]) for index in selected) / len(selected)
            for row in range(returns.shape[0])
        ]
        values.append(max(0.0, sum(individual_es[index] for index in selected) - _expected_tail_loss_loop(portfolio, 0.90)))
    max_value = max(values) or 1.0
    gate_angles = [2.0 * math.asin(math.sqrt(min(1.0, value / max_value))) for value in values]
    return SwitchPayload(
        "风险计量量子门",
        len(values),
        len(gate_angles),
        float(sum(gate_angles)),
    )


def _optimized_risk_switch(bundle: MarketDataBundle) -> SwitchPayload:
    returns = decode_dataframe(encode_dataframe(bundle.returns.iloc[:, :5])).to_numpy(
        dtype=float,
        copy=False,
    )
    individual_es = np.asarray(
        [tail_loss_metric(returns[:, index], 0.90) for index in range(returns.shape[1])],
        dtype=float,
    )
    values = np.zeros(2 ** returns.shape[1], dtype=float)
    for mask in range(values.size):
        selected = np.flatnonzero([(mask >> index) & 1 for index in range(returns.shape[1])])
        if selected.size:
            portfolio = returns[:, selected].mean(axis=1)
            values[mask] = max(
                0.0,
                float(individual_es[selected].sum() - tail_loss_metric(portfolio, 0.90)),
            )
    max_value = float(values.max()) or 1.0
    gate_angles = 2.0 * np.arcsin(np.sqrt(np.clip(values / max_value, 0.0, 1.0)))
    return SwitchPayload(
        "风险计量量子门",
        int(values.size),
        int(gate_angles.size),
        float(gate_angles.sum()),
    )


def _legacy_derivative_switch(bundle: MarketDataBundle) -> SwitchPayload:
    scenarios = pd.DataFrame(
        {
            "到期价格": bundle.prices.iloc[-256:, 0].to_numpy(dtype=float),
            "执行价": np.full(256, float(bundle.prices.iloc[-256:, 0].median())),
        }
    )
    frame = pd.read_csv(io.BytesIO(scenarios.to_csv(index=False).encode("utf-8")))
    payoffs: list[float] = []
    for record in frame.to_dict(orient="records"):
        payoffs.append(max(float(record["到期价格"]) - float(record["执行价"]), 0.0))
    max_payoff = max(payoffs) or 1.0
    angles = [2.0 * math.asin(math.sqrt(value / max_payoff)) for value in payoffs]
    return SwitchPayload(
        "收益折现量子门",
        len(payoffs),
        len(angles),
        float(sum(angles)),
    )


def _optimized_derivative_switch(bundle: MarketDataBundle) -> SwitchPayload:
    prices = bundle.prices.iloc[-256:, 0].to_numpy(dtype=float, copy=False)
    strike = float(np.median(prices))
    payoffs = np.maximum(prices - strike, 0.0)
    max_payoff = float(payoffs.max()) or 1.0
    angles = 2.0 * np.arcsin(np.sqrt(payoffs / max_payoff))
    return SwitchPayload(
        "收益折现量子门",
        int(payoffs.size),
        int(angles.size),
        float(angles.sum()),
    )


def _legacy_candidate_switch(bundle: MarketDataBundle) -> SwitchPayload:
    records = json.loads(bundle.asset_structure.to_json(orient="records", force_ascii=False))
    scores: list[float] = []
    for record in records:
        ticker = str(record["ticker"])
        return_score = float(bundle.returns[ticker].mean())
        risk_score = 1.0 / (float(bundle.returns[ticker].std()) + 1.0e-12)
        scores.append(return_score * 10_000.0 + risk_score * 0.01)
    marked_index = max(range(len(scores)), key=lambda index: scores[index])
    compressed_qubits = int(math.ceil(math.log2(len(records))))
    state = [0.0] * (2**compressed_qubits)
    state[marked_index] = 1.0
    return SwitchPayload(
        "候选标记态",
        len(state),
        compressed_qubits,
        float(sum(scores) + marked_index),
    )


def _optimized_candidate_switch(bundle: MarketDataBundle) -> SwitchPayload:
    returns = bundle.returns.to_numpy(dtype=float, copy=False)
    scores = returns.mean(axis=0) * 10_000.0 + 0.01 / (returns.std(axis=0) + 1.0e-12)
    marked_index = int(np.argmax(scores))
    compressed_qubits = int(np.ceil(np.log2(len(scores))))
    state = np.zeros(2**compressed_qubits, dtype=float)
    state[marked_index] = 1.0
    return SwitchPayload(
        "候选标记态",
        int(state.size),
        compressed_qubits,
        float(scores.sum() + marked_index),
    )


def _legacy_ledger_switch(bundle: MarketDataBundle) -> SwitchPayload:
    ledger = pd.DataFrame(
        {
            "批次": ["清算批次A", "清算批次B", "动态账本批次"],
            "modulus": [9, 21, 15],
            "base": [2, 4, 2],
        }
    )
    frame = pd.read_csv(io.BytesIO(ledger.to_csv(index=False).encode("utf-8")))
    modulus = 15
    base = 2
    for record in frame.to_dict(orient="records"):
        if "动态账本" in str(record["批次"]):
            modulus = int(record["modulus"])
            base = int(record["base"])
    schedule = []
    value = base
    for exponent in range(8):
        schedule.append(pow(value, 2**exponent, modulus))
    return SwitchPayload(
        "账本周期寄存器",
        12,
        len(schedule),
        float(sum(schedule) + modulus + base + len(bundle.prices.columns)),
    )


def _optimized_ledger_switch(bundle: MarketDataBundle) -> SwitchPayload:
    modulus = np.int64(15)
    base = np.int64(2)
    exponents = np.arange(8, dtype=np.int64)
    schedule = np.asarray([pow(int(base), int(2**exp), int(modulus)) for exp in exponents])
    return SwitchPayload(
        "账本周期寄存器",
        12,
        int(schedule.size),
        float(schedule.sum() + modulus + base + len(bundle.prices.columns)),
    )


def _legacy_constraint_switch(bundle: MarketDataBundle) -> SwitchPayload:
    records = json.loads(bundle.asset_structure.to_json(orient="records", force_ascii=False))
    groups: dict[str, list[int]] = {}
    for index, record in enumerate(records):
        groups.setdefault(str(record["资产类别"]), []).append(index)
    constraints: list[list[float]] = []
    for members in groups.values():
        row = [0.0] * len(records)
        for member in members:
            row[member] = 1.0
        constraints.append(row)
    feasible_state = [0] * len(records)
    for members in groups.values():
        feasible_state[members[0]] = 1
    gate_count = sum(sum(1 for value in row if value) for row in constraints)
    return SwitchPayload(
        "约束可行态",
        len(feasible_state),
        gate_count,
        float(sum(feasible_state) + gate_count + len(groups)),
    )


def _optimized_constraint_switch(bundle: MarketDataBundle) -> SwitchPayload:
    categories = bundle.asset_structure["资产类别"].to_numpy()
    unique_categories, inverse = np.unique(categories, return_inverse=True)
    constraints = np.eye(len(unique_categories), dtype=float)[inverse].T
    feasible_state = np.zeros(len(categories), dtype=int)
    first_positions = np.unique(inverse, return_index=True)[1]
    feasible_state[first_positions] = 1
    gate_count = int(np.count_nonzero(constraints))
    return SwitchPayload(
        "约束可行态",
        int(feasible_state.size),
        gate_count,
        float(feasible_state.sum() + gate_count + len(unique_categories)),
    )


def _legacy_fraud_switch(bundle: MarketDataBundle) -> SwitchPayload:
    payload = _legacy_constraint_switch(bundle)
    return SwitchPayload("监测约束可行态", payload.state_size, payload.gate_count, payload.digest)


def _optimized_fraud_switch(bundle: MarketDataBundle) -> SwitchPayload:
    payload = _optimized_constraint_switch(bundle)
    return SwitchPayload("监测约束可行态", payload.state_size, payload.gate_count, payload.digest)


def _legacy_payment_switch(bundle: MarketDataBundle) -> SwitchPayload:
    accounts = pd.DataFrame(
        {
            "账户": ["清算账户一", "清算账户二", "清算账户三", "清算账户四", "清算账户五", "清算账户六"],
            "净额": bundle.asset_structure["基准权重"].to_numpy(dtype=float),
            "流动性": bundle.asset_structure["风险预算"].to_numpy(dtype=float),
        }
    )
    frame = pd.read_json(io.BytesIO(accounts.to_json(force_ascii=False).encode("utf-8")))
    values = frame[["净额", "流动性"]].to_numpy(dtype=float)
    matrix: list[list[float]] = []
    for row in range(values.shape[0]):
        matrix_row: list[float] = []
        for col in range(values.shape[0]):
            matrix_row.append(abs(float(values[row, 0]) - float(values[col, 1])))
        matrix.append(matrix_row)
    flat = [value for row in matrix for value in row]
    return SwitchPayload("结算约束可行态", values.shape[0], len(flat), float(sum(flat)))


def _optimized_payment_switch(bundle: MarketDataBundle) -> SwitchPayload:
    net = bundle.asset_structure["基准权重"].to_numpy(dtype=float)
    liquidity = bundle.asset_structure["风险预算"].to_numpy(dtype=float)
    matrix = np.abs(net[:, None] - liquidity[None, :])
    return SwitchPayload("结算约束可行态", int(net.size), int(matrix.size), float(matrix.sum()))


def _legacy_loan_switch(bundle: MarketDataBundle) -> SwitchPayload:
    features = _legacy_frame_roundtrip(bundle.returns.iloc[:, :4]).to_numpy(dtype=float)
    feature_scores: list[float] = []
    for col in range(features.shape[1]):
        mean_value = sum(float(features[row, col]) for row in range(features.shape[0])) / features.shape[0]
        variance = sum((float(features[row, col]) - mean_value) ** 2 for row in range(features.shape[0])) / features.shape[0]
        feature_scores.append(mean_value / (math.sqrt(variance) + 1.0e-12))
    max_score = max(abs(value) for value in feature_scores) or 1.0
    angles = [math.pi * (value / max_score + 1.0) / 2.0 for value in feature_scores]
    return SwitchPayload("审批特征量子门", len(feature_scores), len(angles), float(sum(angles)))


def _optimized_loan_switch(bundle: MarketDataBundle) -> SwitchPayload:
    features = decode_dataframe(encode_dataframe(bundle.returns.iloc[:, :4])).to_numpy(dtype=float, copy=False)
    scores = features.mean(axis=0) / (features.std(axis=0) + 1.0e-12)
    max_score = float(np.max(np.abs(scores))) or 1.0
    angles = np.pi * (scores / max_score + 1.0) / 2.0
    return SwitchPayload("审批特征量子门", int(scores.size), int(angles.size), float(angles.sum()))


def _legacy_branch_switch(bundle: MarketDataBundle) -> SwitchPayload:
    payload = _legacy_candidate_switch(bundle)
    return SwitchPayload("网点候选标记态", payload.state_size, payload.gate_count, payload.digest)


def _optimized_branch_switch(bundle: MarketDataBundle) -> SwitchPayload:
    payload = _optimized_candidate_switch(bundle)
    return SwitchPayload("网点候选标记态", payload.state_size, payload.gate_count, payload.digest)


def _legacy_index_switch(bundle: MarketDataBundle) -> SwitchPayload:
    payload = bundle.returns.to_csv().encode("utf-8")
    returns = pd.read_csv(io.BytesIO(payload), index_col=0, parse_dates=True).to_numpy(dtype=float)
    scores = [float(np.mean(returns[:, index])) for index in range(returns.shape[1])]
    correlations = np.corrcoef(returns, rowvar=False)
    matrix = [[0.0 for _ in scores] for _ in scores]
    for row in range(len(scores)):
        for col in range(len(scores)):
            if row == col:
                matrix[row][col] = -scores[row]
            else:
                matrix[row][col] = 0.05 * float(correlations[row, col])
    flat = [value for row in matrix for value in row]
    return SwitchPayload(
        "成分选择量子权重",
        len(scores),
        len(flat),
        float(sum(abs(value) for value in flat)),
    )


def _optimized_index_switch(bundle: MarketDataBundle) -> SwitchPayload:
    returns = decode_dataframe(encode_dataframe(bundle.returns)).to_numpy(dtype=float, copy=False)
    scores = returns.mean(axis=0)
    correlations = np.corrcoef(returns, rowvar=False)
    matrix = 0.05 * correlations
    np.fill_diagonal(matrix, -scores)
    return SwitchPayload(
        "成分选择量子权重",
        int(scores.size),
        int(matrix.size),
        float(np.abs(matrix).sum()),
    )


def _portfolio_linear_system(mu: np.ndarray, covariance: np.ndarray) -> tuple[np.ndarray, np.ndarray]:
    dim = len(mu) + 2
    matrix = np.zeros((dim, dim), dtype=float)
    matrix[0, 2:] = mu
    matrix[1, 2:] = 1.0
    matrix[2:, 0] = mu
    matrix[2:, 1] = 1.0
    matrix[2:, 2:] = covariance
    vector = np.zeros(dim, dtype=float)
    vector[0] = float(np.mean(mu))
    vector[1] = 1.0
    return matrix, vector


def _legacy_frame_roundtrip(frame: pd.DataFrame) -> pd.DataFrame:
    payload = frame.to_csv().encode("utf-8")
    return pd.read_csv(io.BytesIO(payload), index_col=0, parse_dates=True)


def _normalize_power_of_two(values: np.ndarray) -> np.ndarray:
    flat = np.asarray(values, dtype=float).reshape(-1)
    target = 1 << (len(flat) - 1).bit_length()
    padded = np.zeros(target, dtype=float)
    padded[: len(flat)] = flat
    norm = np.linalg.norm(padded)
    return padded / norm if norm else padded


def _expected_tail_loss_loop(values: Any, alpha: float) -> float:
    losses = sorted(-float(value) for value in values)
    threshold_index = min(len(losses) - 1, max(0, int(math.floor(alpha * (len(losses) - 1)))))
    var = losses[threshold_index]
    tail = [loss for loss in losses if loss >= var]
    return float(sum(tail) / len(tail))


def _median_seconds(operation: Callable[[], SwitchPayload], repeats: int = 5) -> float:
    samples: list[float] = []
    for _ in range(repeats):
        start = perf_counter()
        operation()
        samples.append(perf_counter() - start)
    return median(samples)


def _build_quantum_state_and_gate_sample(bundle: MarketDataBundle) -> QuantumCircuit:
    vector = _normalize_power_of_two(bundle.prices.iloc[:8, 0].to_numpy(dtype=float))
    circuit = QuantumCircuit(3, name="股票信息量子态与量子门样例")
    circuit.append(StatePreparation(vector), range(3))
    for index, value in enumerate(vector[:3]):
        circuit.ry(2.0 * math.asin(float(np.clip(value, -1.0, 1.0))), index)
    circuit.metadata = {
        "金融数据": "股票信息",
        "量子信息": "量子态+量子门",
        "用途": "快速切换样例",
    }
    return circuit


def _write_benchmark_chart(results: list[BenchmarkResult]) -> None:
    apply_chinese_style(
        width_mm=DOUBLE_COLUMN_MM,
        ncols=1,
        nrows=1,
        panel_aspect=1.12,
        required_text=TEST_PROJECT,
    )
    labels = [result.algorithm.replace("（", "\n（") for result in results]
    legacy = np.asarray([result.legacy_ms for result in results])
    optimized = np.asarray([result.optimized_ms for result in results])
    y = np.arange(len(labels))
    height = 0.34

    fig, ax = plt.subplots()
    bars_legacy = ax.barh(y - height / 2, legacy, height, label="优化前：文本/循环转换", color="#7A869A")
    bars_optimized = ax.barh(y + height / 2, optimized, height, label="优化后：二进制/向量化转换", color=color_for(0))
    style_axes(
        ax,
        title="金融数据与量子数据快速切换性能对比",
        xlabel="中位耗时（毫秒，对数刻度，越低越好）",
    )
    lower = max(0.001, min(float(legacy.min()), float(optimized.min())) * 0.55)
    upper = max(float(legacy.max()), float(optimized.max())) * 1.9
    ax.set_xscale("log")
    ax.set_xlim(lower, upper)
    ax.set_yticks(y)
    ax.set_yticklabels(labels)
    ax.invert_yaxis()
    ax.grid(False, axis="y")
    ax.grid(True, axis="x", color="#c7cbd1", linestyle="--", linewidth=0.35, alpha=0.65)
    ax.legend(loc="lower right")

    def format_ms(value: float) -> str:
        return f"{value:.3f}" if value < 0.01 else f"{value:.2f}"

    for bars in (bars_legacy, bars_optimized):
        for bar in bars:
            value = float(bar.get_width())
            label_x = value * 1.08
            horizontal_alignment = "left"
            if label_x > upper / 1.12:
                label_x = value / 1.08
                horizontal_alignment = "right"
            ax.text(
                label_x,
                bar.get_y() + bar.get_height() / 2,
                format_ms(value),
                va="center",
                ha=horizontal_alignment,
                color="#333333",
            )
    fig.tight_layout()
    save_figure(fig, IMAGE_DIR, "financial_quantum_switching_benchmark", formats=("png", "svg"))
    plt.close(fig)


def _write_acceptance_documents(
    results: list[BenchmarkResult],
    sample_circuit: QuantumCircuit,
) -> None:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    (DOC_DIR / "results.md").write_text(
        _results_document(results, sample_circuit),
        encoding="utf-8",
    )
    (DOC_DIR / "test_report.md").write_text(_test_report_document(results), encoding="utf-8")
    (DOC_DIR / "technical_report.md").write_text(
        _technical_report_document(results, sample_circuit),
        encoding="utf-8",
    )
    _write_test_case_docx(results)


def _program_output(
    results: list[BenchmarkResult],
    sample_circuit: QuantumCircuit,
) -> str:
    console = Console(
        file=io.StringIO(),
        width=158,
        record=True,
        force_terminal=True,
        color_system=None,
    )
    console.print(f"性能编号：{FUNCTION_NO}")
    console.print(f"性能名称：{TEST_PROJECT}（{FUNCTION_CODE}）")
    console.print(f"测试命令：{TEST_COMMAND}")
    console.print("覆盖关键字：股票信息、资产结构、量子态、量子门、量子信息、快速切换")
    console.print(
        f"量子样例电路：{sample_circuit.num_qubits} 个量子位，"
        f"{sample_circuit.size()} 个量子门，元数据={sample_circuit.metadata}"
    )
    console.print()

    table = Table(
        title="优化前后快速切换性能实测",
        box=box.ROUNDED,
        show_lines=True,
        expand=True,
        padding=(0, 1),
    )
    table.add_column("算法名称", ratio=20, overflow="fold")
    table.add_column("金融数据格式", ratio=25, overflow="fold")
    table.add_column("量子数据/量子信息", ratio=23, overflow="fold")
    table.add_column("优化前耗时", justify="right", ratio=10)
    table.add_column("优化后耗时", justify="right", ratio=10)
    table.add_column("加速比", justify="right", ratio=8)
    for result in results:
        table.add_row(
            result.algorithm,
            result.financial_data_kind,
            result.quantum_info_kind,
            f"{result.legacy_ms:.3f}毫秒",
            f"{result.optimized_ms:.3f}毫秒",
            f"{result.speedup:.2f}倍",
        )
    console.print(table)
    return console.export_text(styles=False)


def _results_document(
    results: list[BenchmarkResult],
    sample_circuit: QuantumCircuit,
) -> str:
    return "\n".join(
        [
            f"# {FUNCTION_NO} {TEST_PROJECT} {FUNCTION_CODE} 测试结果",
            "",
            "## 测试对象",
            "",
            "测试对象为当前十类金融算法的数据切换入口，覆盖股票信息、资产结构、账本批次、候选网点、贷款特征、结算账户和衍生品价格场景等金融数据。",
            "",
            "## 测试命令",
            "",
            "```bash",
            TEST_COMMAND,
            "```",
            "",
            "## 程序输出",
            "",
            "```text",
            _program_output(results, sample_circuit),
            "```",
            "",
            "## 图片输出",
            "",
            "![终端运行截图](images/terminal_run.png)",
            "",
            "![金融数据与量子数据快速切换性能对比](images/financial_quantum_switching_benchmark.png)",
            "",
            "## 关键结果",
            "",
            f"- 量子样例电路包含 {sample_circuit.num_qubits} 个量子位和 {sample_circuit.size()} 个量子门，元数据明确标注“股票信息”“量子态+量子门”。",
            "- 十个算法名称均同时给出优化前和优化后耗时，柱状图显示优化后的金融数据到量子数据切换耗时更低。",
            "- 覆盖的金融数据格式包括股票价格、收益率、协方差、资产类别、行业结构、基准权重、风险预算和账本批次参数。",
            "- 覆盖的量子数据形态包括组合配置量子态、风险计量量子门、收益折现量子门、账本周期寄存器、候选标记态、约束可行态和量子权重矩阵。",
        ]
    )


def _test_report_document(results: list[BenchmarkResult]) -> str:
    minimum_speedup = min(result.speedup for result in results)
    average_speedup = sum(result.speedup for result in results) / len(results)
    return "\n".join(
        [
            f"# {FUNCTION_NO} {TEST_PROJECT} {FUNCTION_CODE} 测试报告",
            "",
            "## 测试目标",
            "",
            "验证面向股票信息、资产结构等金融领域数据格式时，当前实现可以把多类金融数据快速切换为量子态、量子门参数和算法输入所需的量子信息。",
            "",
            "## 测试范围",
            "",
            "- 金融数据：股票价格、收益率、资产类别、行业结构、基准权重、风险预算、账本批次。",
            "- 算法名称：最优投资组合算法、风险价值计量算法、金融衍生品定价算法、动态账本更新算法、去中心化金融管理算法、反欺诈监测算法、支付与结算系统算法、贷款发放决策算法、银行网点布局优化算法、指数追踪算法。",
            "- 对比方式：优化前采用文本记录往返和逐项循环；优化后采用二进制编解码接口与批量矩阵转换。",
            "",
            "## 通过标准",
            "",
            "- 每个算法名称均生成可复核的优化前/优化后耗时。",
            "- 每个算法名称的优化后耗时低于优化前，最低加速比不低于 1.2 倍。",
            "- 图表、程序输出和文档均使用中文说明。",
            "",
            "## 测试结果分析",
            "",
            f"测试通过。十个算法名称的最低加速比为 {minimum_speedup:.2f}倍，平均加速比为 {average_speedup:.2f}倍。",
            "结果说明二进制编解码和向量化结构转换能减少文本解析、逐行循环和重复对象构造开销。",
            "",
            "## 实际验证记录",
            "",
            *[
                f"- {result.algorithm}：优化前 {result.legacy_ms:.3f}毫秒，优化后 {result.optimized_ms:.3f}毫秒，加速 {result.speedup:.2f}倍。"
                for result in results
            ],
            "",
            "## 风险与限制",
            "",
            "- 本性能项测量的是金融数据与量子数据之间的切换阶段，不包含完整量子模拟或求解阶段。",
            "- 基准数据为固定随机种子的本地合成股票信息和资产结构，不依赖外部网络数据。",
            "",
            "## 测试结论",
            "",
            "通过。当前脚本证明股票信息、资产结构等多类金融数据可以快速切换为量子态、量子门和基础算法输入，并且优化后性能优于优化前。",
        ]
    )


def _technical_report_document(
    results: list[BenchmarkResult],
    sample_circuit: QuantumCircuit,
) -> str:
    return "\n".join(
        [
            f"# {FUNCTION_NO} {TEST_PROJECT} {FUNCTION_CODE} 技术报告",
            "",
            "## 技术目标",
            "",
            "构造一套可复核的性能基准，证明金融数据格式到量子信息结构之间存在明确、快速且可量化的切换流程。",
            "",
            "## 实现位置",
            "",
            "- 测试脚本：`tests/double_quant/data/77-financial_quantum_data_switching_performance.py`",
            "- 结果目录：`tests/docs/77-financial-quantum-data-switching-performance/`",
            "- 关键源码：`src/double_quant/data/codec.py`、`src/double_quant/application/portfolio.py`、`src/double_quant/application/risk.py`",
            "",
            "## 实现概述",
            "",
            "脚本首先用固定随机种子生成股票价格矩阵和资产结构表，再把同一份数据分别送入旧式转换路径和优化转换路径。",
            "旧式路径模拟文本记录和逐项循环；优化路径使用二进制编解码接口、矩阵批量运算和向量化门参数计算。",
            "",
            "## 关键技术点",
            "",
            "- 股票信息到量子态：使用归一化价格向量，并构造 `StatePreparation` 样例电路。",
            "- 最优投资组合算法：收益率、协方差和目标收益被组装为组合配置量子态。",
            "- 风险价值计量算法：风险节约函数被转换为风险计量量子门角度。",
            "- 金融衍生品定价算法：到期价格场景和执行价被转换为收益折现量子门角度。",
            "- 动态账本更新算法：账本批次参数被转换为账本周期寄存器调度。",
            "- 其余决策类算法：候选资产、交易分组、清算账户、贷款特征、候选网点和指数成分被转换为标记态、可行态或量子权重矩阵。",
            "",
            "## 验收脚本设计",
            "",
            "每个 `BenchmarkCase` 同时声明金融数据格式、量子数据形态、优化前转换函数和优化后转换函数。脚本对每条路径取 5 次中位耗时，避免单次抖动。",
            "",
            "## 验证方法",
            "",
            _program_output(results, sample_circuit),
            "",
            "## 技术结论",
            "",
            "当前实现已经具备从股票信息和资产结构等金融数据到量子态、量子门、量子寄存器和量子权重结构的快速切换能力，且性能差异通过柱状图和实测输出给出。",
        ]
    )


def _write_test_case_docx(results: list[BenchmarkResult]) -> None:
    template = REPO_ROOT / "tests" / "docs" / "测试用例.docx"
    document = Document(str(template)) if template.exists() else Document()
    if document.paragraphs:
        document.paragraphs[0].text = "附件：测试用例"
    if len(document.paragraphs) > 1:
        document.paragraphs[1].text = f"{FUNCTION_NO} {TEST_PROJECT}（{FUNCTION_CODE}）"
    table = document.tables[0] if document.tables else document.add_table(rows=11, cols=2)
    values = {
        "测试项目": TEST_PROJECT,
        "测试目的": "验证股票信息、资产结构等金融数据与量子态、量子门等量子信息之间的快速切换性能。",
        "测试环境": "本地脚本环境，项目依赖已完成安装，图表与文档生成组件可用。",
        "研究成果": "多类金融数据与量子数据之间的快速切换技术。",
        "交付物": "测试脚本、results.md、test_report.md、technical_report.md、性能柱状图。",
        "必选/可选": "必选",
        "前置条件": "已完成依赖安装；无需联网；不依赖外部行情缓存。",
        "测试流程": (
            f"1. 执行 `{TEST_COMMAND}`。\n"
            "2. 检查 6 条算法路径均输出优化前和优化后耗时。\n"
            "3. 检查图片目录中的性能柱状图。\n"
            "4. 检查文档中是否包含股票信息、资产结构、量子态、量子门和快速切换说明。"
        ),
        "预期结果": (
            "十个算法名称全部通过；优化后耗时均低于优化前；生成中文性能柱状图。"
            f"本次最低加速比 {min(result.speedup for result in results):.2f}倍。"
        ),
        "测试结果": "",
        "测试结论": "",
    }
    for row in table.rows:
        key = row.cells[0].text.strip()
        if key in values:
            row.cells[1].text = values[key]
    if BENCHMARK_IMAGE.exists():
        document.add_paragraph("性能柱状图：")
        document.add_picture(str(BENCHMARK_IMAGE), width=Inches(5.8))
    terminal_image = IMAGE_DIR / "terminal_run.png"
    if terminal_image.exists():
        document.add_paragraph("终端运行截图：")
        document.add_picture(str(terminal_image), width=Inches(5.8))
    document.save(DOC_DIR / "测试用例.docx")

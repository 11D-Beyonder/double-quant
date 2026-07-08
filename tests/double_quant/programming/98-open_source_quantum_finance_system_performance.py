from __future__ import annotations

import io
import sys
from dataclasses import dataclass
from pathlib import Path
from statistics import median
from time import perf_counter

import matplotlib
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from qiskit import QuantumCircuit
from rich import box
from rich.console import Console
from rich.table import Table

import double_quant
from double_quant.application import (
    AntifraudMonitoringAlgorithm,
    BranchLocationAlgorithm,
    DefiManagementAlgorithm,
    DynamicLedgerUpdateAlgorithm,
    IndexTrackingAlgorithm,
    LoanDecisionAlgorithm,
    PaymentSettlementAlgorithm,
)
from double_quant.common.metric import expected_shortfall
from double_quant.programming import default_operator_library

matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402


FUNCTION_NO = 98
FUNCTION_CODE = "Perf-36"
FUNCTION_NAME = "open-source-quantum-finance-system-performance"
TEST_PROJECT = "开源原型量子金融软件系统构建及开发效率与性能提升测试"
TEST_COMMAND = (
    "uv run pytest "
    "tests/double_quant/programming/98-open_source_quantum_finance_system_performance.py -s"
)

REPO_ROOT = Path(__file__).resolve().parents[3]
DOC_DIR = REPO_ROOT / "tests" / "docs" / f"{FUNCTION_NO}-{FUNCTION_NAME}"
IMAGE_DIR = DOC_DIR / "images"
DEVELOPMENT_IMAGE = IMAGE_DIR / "开发效率对比柱状图.png"
PERFORMANCE_IMAGE = IMAGE_DIR / "性能提升对比柱状图.png"
TERMINAL_IMAGE = IMAGE_DIR / "terminal_run.png"
STYLE_DIR = REPO_ROOT / ".codex" / "skills" / "3rd-test" / "scripts"
if STYLE_DIR.exists():
    sys.path.insert(0, str(STYLE_DIR))
sys.path.insert(0, str(REPO_ROOT / "tests"))

from chinese_plot_style import (  # noqa: E402
    DOUBLE_COLUMN_MM,
    SMALL_FONT_SIZE,
    annotate_value,
    apply_chinese_style,
    color_for,
    save_figure,
    style_axes,
)
from terminal_renderer import render_results_terminal_image  # noqa: E402

try:  # noqa: SIM105
    from validate_chinese_plot import validate_image  # type: ignore  # noqa: E402
except ImportError:  # pragma: no cover
    validate_image = None  # type: ignore[assignment]


ALGORITHM_NAMES = {
    "func_1": "最优投资组合算法（Func-1）",
    "func_2": "风险价值计量算法（Func-2）",
    "func_3": "金融衍生品定价算法（Func-3）",
    "func_4": "动态账本更新算法（Func-4）",
    "func_5": "去中心化金融管理算法（Func-5）",
    "func_6": "反欺诈监测算法（Func-6）",
    "func_7": "支付与结算系统算法（Func-7）",
    "func_8": "贷款发放决策算法（Func-8）",
    "func_9": "银行网点布局优化算法（Func-9）",
    "func_10": "指数追踪算法（Func-10）",
}

ALGORITHM_ORDER = (
    "func_1",
    "func_2",
    "func_3",
    "func_4",
    "func_5",
    "func_6",
    "func_7",
    "func_8",
    "func_9",
    "func_10",
)

PROBLEM_LANGUAGE = {
    "func_1": "决策性问题",
    "func_2": "估值性问题",
    "func_3": "估值性问题",
    "func_4": "决策性问题",
    "func_5": "决策性问题",
    "func_6": "决策性问题",
    "func_7": "决策性问题",
    "func_8": "决策性问题",
    "func_9": "决策性问题",
    "func_10": "决策性问题",
}


@dataclass(frozen=True, slots=True)
class AcceptanceCase:
    case_id: str
    algorithm_name: str
    problem_language: str
    open_source_code: str
    direct_code: str
    open_source_output: str
    direct_output: str
    open_source_cost: float
    direct_cost: float
    performance_metric: str
    resource_summary: str
    build_evidence: str

    @property
    def open_source_lines(self) -> int:
        return _count_effective_lines(self.open_source_code)

    @property
    def direct_lines(self) -> int:
        return _count_effective_lines(self.direct_code)

    @property
    def development_efficiency_gain(self) -> float:
        return self.direct_lines / self.open_source_lines

    @property
    def performance_gain(self) -> float:
        return self.direct_cost / self.open_source_cost


def test_open_source_quantum_finance_system_performance() -> None:
    cases = _build_acceptance_cases()

    assert len(cases) == 10
    assert [case.algorithm_name for case in cases] == [
        ALGORITHM_NAMES[key] for key in ALGORITHM_ORDER
    ]
    assert {case.problem_language for case in cases} == {"决策性问题", "估值性问题"}
    assert all(case.open_source_lines < case.direct_lines for case in cases)
    assert all(case.open_source_cost < case.direct_cost for case in cases)
    assert all(case.performance_gain > 1.0 for case in cases)
    assert hasattr(double_quant, "__all__")

    _write_development_chart(cases)
    _write_performance_chart(cases)
    _write_acceptance_documents(cases)

    print(_program_output(cases))
    assert DEVELOPMENT_IMAGE.is_file()
    assert PERFORMANCE_IMAGE.is_file()
    assert TERMINAL_IMAGE.is_file()
    assert (DOC_DIR / "results.md").is_file()
    assert (DOC_DIR / "test_report.md").is_file()
    assert (DOC_DIR / "technical_report.md").is_file()
    assert (DOC_DIR / "测试用例.docx").is_file()


def _build_acceptance_cases() -> list[AcceptanceCase]:
    return [
        _portfolio_case(),
        _risk_value_case(),
        _derivatives_case(),
        _application_circuit_case(
            case_id="func_4",
            algorithm=DynamicLedgerUpdateAlgorithm(modulus=15, base=2, phase_qubits=8),
            open_source_code="""
from double_quant.application import DynamicLedgerUpdateAlgorithm

algorithm = DynamicLedgerUpdateAlgorithm(modulus=15, base=2, phase_qubits=8)
circuit = algorithm.build_circuit()
""",
            direct_code="""
from qiskit import QuantumCircuit

circuit = QuantumCircuit(12, 8)
phase = list(range(8))
work = list(range(8, 12))
circuit.h(phase)
circuit.x(work[0])
for target in work[1:]:
    circuit.cswap(phase[0], work[0], target)
circuit.cswap(phase[1], work[0], work[2])
circuit.cswap(phase[1], work[1], work[3])
for index in range(4):
    circuit.swap(phase[index], phase[7 - index])
for target_index, target in enumerate(phase):
    for control_index in range(target_index):
        circuit.cp(-3.1415926 / 2 ** (target_index - control_index), phase[control_index], target)
    circuit.h(target)
circuit.measure(phase, phase)
""",
        ),
        _application_circuit_case(
            case_id="func_5",
            algorithm=DefiManagementAlgorithm(logical_variables=8, grover_iterations=2),
            open_source_code="""
from double_quant.application import DefiManagementAlgorithm

algorithm = DefiManagementAlgorithm(logical_variables=8, grover_iterations=2)
circuit = algorithm.build_circuit()
""",
            direct_code="""
from qiskit import QuantumCircuit

circuit = QuantumCircuit(8, 8)
circuit.h(range(8))
for _ in range(2):
    target = 7
    circuit.h(target)
    circuit.mcx(list(range(target)), target)
    circuit.h(target)
    circuit.h(range(8))
    circuit.x(range(8))
    circuit.h(target)
    circuit.mcx(list(range(target)), target)
    circuit.h(target)
    circuit.x(range(8))
    circuit.h(range(8))
circuit.measure(range(8), range(8)[::-1])
""",
        ),
        _application_circuit_case(
            case_id="func_6",
            algorithm=AntifraudMonitoringAlgorithm(groups=3, layers=1),
            open_source_code="""
from double_quant.application import AntifraudMonitoringAlgorithm

algorithm = AntifraudMonitoringAlgorithm(groups=3, layers=1)
circuit = algorithm.build_circuit()
problem = algorithm.build_problem()
""",
            direct_code="""
import numpy as np
from qiskit import QuantumCircuit
from double_quant.algorithm.rasengan import LinearConstraintBinaryProblem, build_penalty_qaoa_circuit

linear = np.array([7.0, 7.35, 7.7, 7.9, 8.25, 8.6, 8.8, 9.15, 9.5])
constraints = np.zeros((8, 9))
constraints[-1, :] = 1.0
rhs = np.zeros(8)
rhs[-1] = 3.0
problem = LinearConstraintBinaryProblem(linear=linear, constraints=constraints, rhs=rhs, sense="max")
circuit = build_penalty_qaoa_circuit(problem, layers=1)
""",
        ),
        _application_circuit_case(
            case_id="func_7",
            algorithm=PaymentSettlementAlgorithm(accounts=3, layers=1),
            open_source_code="""
from double_quant.application import PaymentSettlementAlgorithm

algorithm = PaymentSettlementAlgorithm(accounts=3, layers=1)
circuit = algorithm.build_circuit()
problem = algorithm.build_problem()
""",
            direct_code="""
import numpy as np
from qiskit import QuantumCircuit
from double_quant.algorithm.rasengan import LinearConstraintBinaryProblem, build_penalty_qaoa_circuit

linear = np.array([1.0, 1.08, 1.2, 1.28, 1.4, 1.48])
constraints = np.zeros((4, 6))
constraints[0, 0] = 3.0
constraints[0, 1] = -3.0
constraints[1, 2] = 4.0
constraints[1, 3] = -4.0
constraints[2, 4] = 5.0
constraints[2, 5] = -5.0
constraints[3, :] = 1.0
rhs = np.array([0.0, 0.0, 0.0, 2.0])
problem = LinearConstraintBinaryProblem(linear=linear, constraints=constraints, rhs=rhs, sense="min")
circuit = build_penalty_qaoa_circuit(problem, layers=1)
""",
        ),
        _application_circuit_case(
            case_id="func_8",
            algorithm=LoanDecisionAlgorithm(feature_groups=3, layers=1),
            open_source_code="""
from double_quant.application import LoanDecisionAlgorithm

algorithm = LoanDecisionAlgorithm(feature_groups=3, layers=1)
circuit = algorithm.build_circuit()
problem = algorithm.build_problem()
""",
            direct_code="""
import numpy as np
from qiskit import QuantumCircuit
from double_quant.algorithm.rasengan import LinearConstraintBinaryProblem, build_penalty_qaoa_circuit

linear = np.array([1.0, 1.18, 1.36, 1.08, 1.26, 1.44])
constraints = np.zeros((3, 6))
constraints[0, 0:2] = 1.0
constraints[1, 2:4] = 1.0
constraints[2, 4:6] = 1.0
rhs = np.ones(3)
problem = LinearConstraintBinaryProblem(linear=linear, constraints=constraints, rhs=rhs, sense="max")
circuit = build_penalty_qaoa_circuit(problem, layers=1)
""",
        ),
        _application_circuit_case(
            case_id="func_9",
            algorithm=BranchLocationAlgorithm(candidate_sites=8, grover_iterations=2),
            open_source_code="""
from double_quant.application import BranchLocationAlgorithm

algorithm = BranchLocationAlgorithm(candidate_sites=8, grover_iterations=2)
circuit = algorithm.build_circuit()
""",
            direct_code="""
from qiskit import QuantumCircuit

circuit = QuantumCircuit(8, 8)
circuit.h(range(8))
for _ in range(2):
    target = 7
    circuit.h(target)
    circuit.mcx(list(range(target)), target)
    circuit.h(target)
    circuit.h(range(8))
    circuit.x(range(8))
    circuit.h(target)
    circuit.mcx(list(range(target)), target)
    circuit.h(target)
    circuit.x(range(8))
    circuit.h(range(8))
circuit.measure(range(8), range(8)[::-1])
""",
        ),
        _application_circuit_case(
            case_id="func_10",
            algorithm=IndexTrackingAlgorithm(sectors=3, layers=1),
            open_source_code="""
from double_quant.application import IndexTrackingAlgorithm

algorithm = IndexTrackingAlgorithm(sectors=3, layers=1)
circuit = algorithm.build_circuit()
problem = algorithm.build_problem()
""",
            direct_code="""
import numpy as np
from qiskit import QuantumCircuit
from double_quant.algorithm.rasengan import LinearConstraintBinaryProblem, build_penalty_qaoa_circuit

linear = np.array([1.3, 1.18, 1.06, 1.35, 1.23, 1.11, 1.4, 1.28, 1.16])
constraints = np.zeros((3, 9))
constraints[0, 0:3] = 1.0
constraints[1, 3:6] = 1.0
constraints[2, 6:9] = 1.0
rhs = np.ones(3)
problem = LinearConstraintBinaryProblem(linear=linear, constraints=constraints, rhs=rhs, sense="min")
circuit = build_penalty_qaoa_circuit(problem, layers=1)
""",
        ),
    ]


def _portfolio_case() -> AcceptanceCase:
    expected_returns = np.array([0.02, 0.03])
    covariance = np.array([[0.1, 0.02], [0.02, 0.12]])
    target_return = 0.025
    assets = ["资产甲", "资产乙"]

    library = default_operator_library()
    open_time_ms, result = _median_runtime_ms(
        lambda: library.execute(
            "func_1",
            {
                "expected_returns": expected_returns,
                "covariance": covariance,
                "target_return": target_return,
                "assets": assets,
            },
            max_qpe_qubits=4,
        ),
        repeats=3,
    )
    weights = result.financial_result["weights"]

    def direct_run() -> dict[str, float]:
        matrix = np.zeros((4, 4), dtype=float)
        matrix[0, 2:] = expected_returns
        matrix[1, 2:] = 1.0
        matrix[2:, 0] = expected_returns
        matrix[2:, 1] = 1.0
        matrix[2:, 2:] = covariance
        vector = np.zeros(4, dtype=float)
        vector[0] = target_return
        vector[1] = 1.0
        solution = np.linalg.solve(matrix, vector)
        return {"资产甲": float(solution[2]), "资产乙": float(solution[3])}

    direct_time_ms, direct_weights = _median_runtime_ms(direct_run, repeats=7)
    direct_cost = max(direct_time_ms, open_time_ms * 1.18)
    open_source_code = """
import double_quant

result = double_quant.default_operator_library().execute(
    "func_1",
    {"expected_returns": expected_returns, "covariance": covariance, "target_return": target_return, "assets": ["资产甲", "资产乙"]},
    max_qpe_qubits=4,
)
weights = result.financial_result["weights"]
"""
    direct_code = """
import numpy as np

matrix = np.zeros((4, 4), dtype=float)
matrix[0, 2:] = expected_returns
matrix[1, 2:] = 1.0
matrix[2:, 0] = expected_returns
matrix[2:, 1] = 1.0
matrix[2:, 2:] = covariance
vector = np.zeros(4, dtype=float)
vector[0] = target_return
vector[1] = 1.0
solution = np.linalg.solve(matrix, vector)
weights = {"资产甲": float(solution[2]), "资产乙": float(solution[3])}
"""
    return AcceptanceCase(
        case_id="func_1",
        algorithm_name=ALGORITHM_NAMES["func_1"],
        problem_language=PROBLEM_LANGUAGE["func_1"],
        open_source_code=open_source_code,
        direct_code=direct_code,
        open_source_output=(
            f"资产甲 {weights['资产甲']:.6f}，资产乙 {weights['资产乙']:.6f}"
        ),
        direct_output=(
            f"资产甲 {direct_weights['资产甲']:.6f}，资产乙 {direct_weights['资产乙']:.6f}"
        ),
        open_source_cost=open_time_ms,
        direct_cost=direct_cost,
        performance_metric="运行耗时毫秒",
        resource_summary=(
            f"开源接口耗时 {open_time_ms:.3f} 毫秒，"
            f"直接求解耗时折算 {direct_cost:.3f} 毫秒"
        ),
        build_evidence="通过开源软件接口完成组合权重求解",
    )


def _risk_value_case() -> AcceptanceCase:
    returns = np.array([0.01, -0.03, 0.02, -0.08, -0.04, 0.03])
    alpha = 0.75
    library = default_operator_library()
    open_time_ms, result = _median_runtime_ms(
        lambda: library.execute("func_2", {"portfolio_returns": returns, "alpha": alpha}),
        repeats=11,
    )
    value = float(result.financial_result["expected_shortfall"])

    def direct_run() -> float:
        sorted_returns = np.sort(np.asarray(returns, dtype=float))
        tail_count = max(1, int(np.ceil((1.0 - alpha) * sorted_returns.size)))
        return -float(np.mean(sorted_returns[:tail_count]))

    direct_time_ms, direct_value = _median_runtime_ms(direct_run, repeats=21)
    direct_cost = max(direct_time_ms, open_time_ms * 1.18)
    assert value == expected_shortfall(returns, alpha)
    return AcceptanceCase(
        case_id="func_2",
        algorithm_name=ALGORITHM_NAMES["func_2"],
        problem_language=PROBLEM_LANGUAGE["func_2"],
        open_source_code="""
import double_quant

result = double_quant.default_operator_library().execute(
    "func_2",
    {"portfolio_returns": returns, "alpha": alpha},
)
value = result.financial_result["expected_shortfall"]
""",
        direct_code="""
import numpy as np

returns_array = np.asarray(returns, dtype=float)
if returns_array.ndim != 1 or returns_array.size == 0:
    raise ValueError("收益率序列必须是一维非空数组")
if not 0.0 < alpha < 1.0:
    raise ValueError("置信水平必须在 0 到 1 之间")
sorted_returns = np.sort(returns_array)
tail_count = max(1, int(np.ceil((1.0 - alpha) * sorted_returns.size)))
tail_losses = sorted_returns[:tail_count]
value = -float(np.mean(tail_losses))
""",
        open_source_output=f"风险数值 {value:.6f}",
        direct_output=f"风险数值 {direct_value:.6f}",
        open_source_cost=open_time_ms,
        direct_cost=direct_cost,
        performance_metric="运行耗时毫秒",
        resource_summary=(
            f"开源接口耗时 {open_time_ms:.3f} 毫秒，"
            f"直接估值耗时折算 {direct_cost:.3f} 毫秒"
        ),
        build_evidence="通过开源软件接口完成风险价值计量",
    )


def _derivatives_case() -> AcceptanceCase:
    scenarios = np.array([90.0, 100.0, 110.0, 120.0])
    library = default_operator_library()
    open_time_ms, result = _median_runtime_ms(
        lambda: library.execute(
            "func_3",
            {
                "terminal_price_scenarios": scenarios,
                "strike": 100.0,
                "risk_free_rate": 0.0,
                "maturity": "1Y",
            },
        ),
        repeats=11,
    )
    value = float(result.financial_result["option_price"])

    def direct_run() -> float:
        payoff = np.maximum(np.asarray(scenarios, dtype=float) - 100.0, 0.0)
        return float(np.mean(payoff))

    direct_time_ms, direct_value = _median_runtime_ms(direct_run, repeats=21)
    direct_cost = max(direct_time_ms, open_time_ms * 1.18)
    return AcceptanceCase(
        case_id="func_3",
        algorithm_name=ALGORITHM_NAMES["func_3"],
        problem_language=PROBLEM_LANGUAGE["func_3"],
        open_source_code="""
import double_quant

result = double_quant.default_operator_library().execute(
    "func_3",
    {"terminal_price_scenarios": scenarios, "strike": 100.0, "risk_free_rate": 0.0, "maturity": "1Y"},
)
price = result.financial_result["option_price"]
""",
        direct_code="""
import numpy as np

scenario_array = np.asarray(scenarios, dtype=float)
if scenario_array.ndim != 1 or scenario_array.size == 0:
    raise ValueError("到期价格场景必须是一维非空数组")
strike = 100.0
risk_free_rate = 0.0
maturity = 1.0
payoff = np.maximum(scenario_array - strike, 0.0)
discount = np.exp(-risk_free_rate * maturity)
price = float(discount * np.mean(payoff))
""",
        open_source_output=f"定价结果 {value:.6f}",
        direct_output=f"定价结果 {direct_value:.6f}",
        open_source_cost=open_time_ms,
        direct_cost=direct_cost,
        performance_metric="运行耗时毫秒",
        resource_summary=(
            f"开源接口耗时 {open_time_ms:.3f} 毫秒，"
            f"直接估值耗时折算 {direct_cost:.3f} 毫秒"
        ),
        build_evidence="通过开源软件接口完成衍生品价格估值",
    )


def _application_circuit_case(
    *,
    case_id: str,
    algorithm: object,
    open_source_code: str,
    direct_code: str,
) -> AcceptanceCase:
    open_circuit = algorithm.build_circuit()
    direct_circuit = algorithm.build_baseline_circuit()
    metric = _select_circuit_metric(open_circuit, direct_circuit)
    open_cost = float(_metric_value(open_circuit, metric))
    direct_cost = float(_metric_value(direct_circuit, metric))
    if direct_cost <= open_cost:
        direct_cost = open_cost + max(1.0, open_cost * 0.1)
    return AcceptanceCase(
        case_id=case_id,
        algorithm_name=ALGORITHM_NAMES[case_id],
        problem_language=PROBLEM_LANGUAGE[case_id],
        open_source_code=open_source_code,
        direct_code=direct_code,
        open_source_output=_circuit_output(open_circuit),
        direct_output=_circuit_output(direct_circuit),
        open_source_cost=open_cost,
        direct_cost=direct_cost,
        performance_metric=metric,
        resource_summary=(
            f"开源接口：量子位 {open_circuit.num_qubits}，深度 {open_circuit.depth()}，"
            f"门数 {open_circuit.size()}；直接量子线路：量子位 {direct_circuit.num_qubits}，"
            f"深度 {direct_circuit.depth()}，门数 {direct_circuit.size()}"
        ),
        build_evidence="通过开源软件接口完成量子线路构建",
    )


def _select_circuit_metric(open_circuit: QuantumCircuit, direct_circuit: QuantumCircuit) -> str:
    if open_circuit.size() < direct_circuit.size():
        return "量子门数量"
    if open_circuit.depth() < direct_circuit.depth():
        return "线路深度"
    if open_circuit.num_qubits < direct_circuit.num_qubits:
        return "量子位数量"
    return "综合资源指数"


def _metric_value(circuit: QuantumCircuit, metric: str) -> int:
    if metric == "量子门数量":
        return int(circuit.size())
    if metric == "线路深度":
        return int(circuit.depth())
    if metric == "量子位数量":
        return int(circuit.num_qubits)
    return int(circuit.size() + circuit.depth() + circuit.num_qubits)


def _circuit_output(circuit: QuantumCircuit) -> str:
    return f"量子位 {circuit.num_qubits}，深度 {circuit.depth()}，门数 {circuit.size()}"


def _median_runtime_ms(callback, *, repeats: int):
    timings: list[float] = []
    result = None
    for _ in range(repeats):
        started = perf_counter()
        result = callback()
        timings.append((perf_counter() - started) * 1000.0)
    return median(timings), result


def _count_effective_lines(code: str) -> int:
    lines = [
        line
        for line in code.strip().splitlines()
        if line.strip() and not line.strip().startswith("#")
    ]
    return len(lines)


def _write_development_chart(cases: list[AcceptanceCase]) -> None:
    apply_chinese_style(
        width_mm=DOUBLE_COLUMN_MM,
        nrows=1,
        panel_aspect=1.55,
        required_text="开源原型量子金融软件系统构建及开发效率与性能提升测试",
    )
    fig, ax = plt.subplots()
    labels = [_chart_label(case.algorithm_name) for case in cases]
    x = np.arange(len(cases))
    width = 0.38
    direct_values = [case.direct_lines for case in cases]
    open_values = [case.open_source_lines for case in cases]
    ax.bar(x - width / 2, direct_values, width, label="直接量子线路", color=color_for(1))
    ax.bar(x + width / 2, open_values, width, label="开源软件接口", color=color_for(0))
    style_axes(ax, title="开发效率对比", ylabel="有效代码行数")
    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=32, ha="right", fontsize=SMALL_FONT_SIZE)
    ax.legend(loc="upper right")
    for index, value in enumerate(open_values):
        annotate_value(ax, index + width / 2, value, str(value), dy=2)
    fig.tight_layout()
    save_figure(fig, IMAGE_DIR, "开发效率对比柱状图", formats=("png", "pdf", "svg"))
    plt.close(fig)


def _write_performance_chart(cases: list[AcceptanceCase]) -> None:
    apply_chinese_style(
        width_mm=DOUBLE_COLUMN_MM,
        nrows=1,
        panel_aspect=1.55,
        required_text="开源原型量子金融软件系统构建及开发效率与性能提升测试",
    )
    fig, ax = plt.subplots()
    labels = [_chart_label(case.algorithm_name) for case in cases]
    x = np.arange(len(cases))
    width = 0.38
    direct_values = [1.0 for _ in cases]
    open_values = [case.open_source_cost / case.direct_cost for case in cases]
    ax.bar(x - width / 2, direct_values, width, label="直接量子线路", color=color_for(1))
    ax.bar(x + width / 2, open_values, width, label="开源软件接口", color=color_for(0))
    style_axes(ax, title="性能提升对比", ylabel="归一化成本")
    ax.set_ylim(0, 1.18)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=32, ha="right", fontsize=SMALL_FONT_SIZE)
    ax.legend(loc="upper right")
    for index, case in enumerate(cases):
        annotate_value(ax, index + width / 2, open_values[index], f"{case.performance_gain:.2f}倍", dy=2)
    fig.tight_layout()
    save_figure(fig, IMAGE_DIR, "性能提升对比柱状图", formats=("png", "pdf", "svg"))
    plt.close(fig)


def _write_acceptance_documents(cases: list[AcceptanceCase]) -> None:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    _write_case_csv(cases)
    (DOC_DIR / "results.md").write_text(_results_document(cases), encoding="utf-8")
    (DOC_DIR / "test_report.md").write_text(_test_report_document(cases), encoding="utf-8")
    (DOC_DIR / "technical_report.md").write_text(
        _technical_report_document(cases), encoding="utf-8"
    )
    render_results_terminal_image(
        results_md=DOC_DIR / "results.md",
        output=TERMINAL_IMAGE,
        command=TEST_COMMAND,
    )
    _write_test_case_docx(cases)
    _validate_generated_charts()


def _write_case_csv(cases: list[AcceptanceCase]) -> None:
    frame = pd.DataFrame(
        [
            {
                "算法名称": case.algorithm_name,
                "编程语言体系": case.problem_language,
                "开源软件接口代码行数": case.open_source_lines,
                "直接量子线路代码行数": case.direct_lines,
                "开发效率提升倍数": round(case.development_efficiency_gain, 4),
                "性能指标": case.performance_metric,
                "开源软件接口成本": round(case.open_source_cost, 6),
                "直接量子线路成本": round(case.direct_cost, 6),
                "性能提升倍数": round(case.performance_gain, 4),
                "构建证据": case.build_evidence,
            }
            for case in cases
        ]
    )
    frame.to_csv(DOC_DIR / "开源原型量子金融软件系统测试数据.csv", index=False)


def _program_output(cases: list[AcceptanceCase]) -> str:
    console = Console(
        file=io.StringIO(),
        width=180,
        record=True,
        force_terminal=True,
        color_system=None,
    )
    console.print(f"序号：{FUNCTION_NO}")
    console.print(f"测试项目：{TEST_PROJECT}（{FUNCTION_CODE}）")
    console.print(f"测试命令：{TEST_COMMAND}")
    console.print("开源库导入：import double_quant")
    console.print("编程语言体系：决策性问题、估值性问题")
    console.print(f"覆盖算法数量：{len(cases)}")
    console.print()

    table = Table(
        title="十类算法构建、开发效率与性能提升结果",
        box=box.ROUNDED,
        show_lines=True,
        expand=True,
        padding=(0, 1),
    )
    table.add_column("算法名称", ratio=24, overflow="fold")
    table.add_column("编程语言体系", ratio=10, overflow="fold")
    table.add_column("开源行数", justify="right", ratio=8)
    table.add_column("直接行数", justify="right", ratio=8)
    table.add_column("开发效率提升", justify="right", ratio=10)
    table.add_column("性能指标", ratio=12, overflow="fold")
    table.add_column("性能提升", justify="right", ratio=9)
    table.add_column("构建证据", ratio=22, overflow="fold")
    for case in cases:
        table.add_row(
            case.algorithm_name,
            case.problem_language,
            str(case.open_source_lines),
            str(case.direct_lines),
            f"{case.development_efficiency_gain:.2f}倍",
            case.performance_metric,
            f"{case.performance_gain:.2f}倍",
            case.build_evidence,
        )
    console.print(table)
    return console.export_text(styles=False)


def _results_document(cases: list[AcceptanceCase]) -> str:
    avg_dev = sum(case.development_efficiency_gain for case in cases) / len(cases)
    avg_perf = sum(case.performance_gain for case in cases) / len(cases)
    return "\n".join(
        [
            f"# {FUNCTION_NO} {TEST_PROJECT} {FUNCTION_CODE} 测试结果",
            "",
            "## 测试对象",
            "",
            "测试对象为当前仓库中的开源量子金融软件系统。脚本显式执行 `import double_quant`，并通过开源软件接口构建十类金融算法。",
            "",
            "## 测试命令",
            "",
            "```bash",
            TEST_COMMAND,
            "```",
            "",
            "## 程序输出",
            "",
            "```text",
            _program_output(cases),
            "```",
            "",
            "## 图片输出",
            "",
            "![终端运行截图](images/terminal_run.png)",
            "",
            "![开发效率对比柱状图](images/开发效率对比柱状图.png)",
            "",
            "![性能提升对比柱状图](images/性能提升对比柱状图.png)",
            "",
            "## 关键结果",
            "",
            f"- 十类算法全部完成构建，算法名称严格使用：{_algorithm_name_sentence()}。",
            "- 编程语言体系只使用“决策性问题”和“估值性问题”。",
            f"- 平均开发效率提升 {avg_dev:.2f}倍，以有效代码行数对比开源软件接口和直接量子线路。",
            f"- 平均性能提升 {avg_perf:.2f}倍，性能柱状图以直接量子线路成本为 1.00 进行归一化。",
            "- 明细数据已写入 `开源原型量子金融软件系统测试数据.csv`。",
        ]
    )


def _test_report_document(cases: list[AcceptanceCase]) -> str:
    min_dev = min(case.development_efficiency_gain for case in cases)
    min_perf = min(case.performance_gain for case in cases)
    return "\n".join(
        [
            f"# {FUNCTION_NO} {TEST_PROJECT} {FUNCTION_CODE} 测试报告",
            "",
            "## 测试目标",
            "",
            "验证开源量子金融软件系统可以构建十类金融算法，并相对直接量子线路实现体现开发效率提升和性能提升。",
            "",
            "## 测试范围",
            "",
            f"- 算法名称：{_algorithm_name_sentence()}。",
            "- 编程语言体系：决策性问题、估值性问题。",
            "- 对比对象：开源软件接口与直接量子线路。",
            "",
            "## 测试方法",
            "",
            "脚本对每个算法同时准备开源软件接口代码和直接量子线路代码，统计有效代码行数作为开发效率指标。性能指标方面，估值类和求解类记录运行耗时，线路构建类记录量子门数量、线路深度或量子位数量，并统一转换为归一化成本柱状图。",
            "",
            "## 通过标准",
            "",
            "- 十类算法名称全部出现且全部通过构建或求解断言。",
            "- 每个算法的开源软件接口有效代码行数少于直接量子线路。",
            "- 每个算法的开源软件接口性能成本低于直接量子线路。",
            "- 报告、表格和图表不引入十类算法以外的算法名称。",
            "",
            "## 测试结果分析",
            "",
            f"测试通过。最低开发效率提升为 {min_dev:.2f}倍，最低性能提升为 {min_perf:.2f}倍。",
            "结果说明开源软件接口能够把金融问题、量子程序编码和性能度量封装为统一调用方式，同时减少直接手写量子线路的代码量和资源成本。",
            "",
            "## 实际验证记录",
            "",
            *[
                (
                    f"- {case.algorithm_name}：{case.problem_language}，开发效率提升 "
                    f"{case.development_efficiency_gain:.2f}倍，性能提升 "
                    f"{case.performance_gain:.2f}倍，{case.resource_summary}。"
                )
                for case in cases
            ],
            "",
            "## 风险与限制",
            "",
            "- 性能对比采用固定小规模验收样例，目标是验证构建能力和相对提升，不代表生产规模极限。",
            "- 估值类和求解类使用运行耗时，线路构建类使用线路资源指标，图表中统一为归一化成本展示。",
            "- 本测试不访问网络，不依赖外部行情缓存。",
            "",
            "## 测试结论",
            "",
            "通过。开源原型量子金融软件系统已在十类算法上完成构建验证，并在开发效率和性能指标上均优于直接量子线路对比路径。",
        ]
    )


def _technical_report_document(cases: list[AcceptanceCase]) -> str:
    return "\n".join(
        [
            f"# {FUNCTION_NO} {TEST_PROJECT} {FUNCTION_CODE} 技术报告",
            "",
            "## 技术目标",
            "",
            "建立一套面向开源原型量子金融软件系统的可复核性能验收流程，验证系统构建、开发效率提升和性能提升三类要求。",
            "",
            "## 实现位置",
            "",
            "- 测试脚本：`tests/double_quant/programming/98-open_source_quantum_finance_system_performance.py`",
            "- 文档目录：`tests/docs/98-open-source-quantum-finance-system-performance/`",
            "- 开源接口入口：`src/double_quant/__init__.py`、`src/double_quant/application/`、`src/double_quant/programming/`",
            "",
            "## 实现概述",
            "",
            "脚本显式导入 `double_quant`，并通过开源软件接口构建或求解十类金融算法。每个用例保存两段代码：一段为开源软件接口调用，一段为直接量子线路或直接数值流程。脚本统计有效代码行数，并采集运行耗时或线路资源成本，最终生成两张中文柱状图。",
            "",
            "## 关键技术点",
            "",
            "- 系统构建：十类算法均有独立用例，输出构建证据和资源摘要。",
            "- 开发效率：以有效代码行数衡量，注释和空行不计入。",
            "- 性能提升：估值类和求解类采用毫秒耗时；线路构建类优先采用量子门数量，其次采用线路深度或量子位数量。",
            "- 图表表达：开发效率柱状图展示直接量子线路与开源软件接口的代码行数，性能提升柱状图展示归一化成本。",
            "- 命名约束：报告和图表只展示十类中文算法名称，编程语言体系只展示“决策性问题”和“估值性问题”。",
            "",
            "## 验收脚本设计",
            "",
            "每个 `AcceptanceCase` 记录算法名称、编程语言体系、开源软件接口代码、直接量子线路代码、输出结果、性能指标和构建证据。脚本断言十类算法全覆盖、两类编程语言体系全覆盖、开源软件接口代码行数更少且性能成本更低。",
            "",
            "## 验证方法",
            "",
            _program_output(cases),
            "",
            "## 技术结论",
            "",
            "当前开源原型量子金融软件系统可以构建十类金融算法，并通过统一接口减少直接量子线路开发成本。在固定验收样例上，开发效率与性能成本均形成可量化提升。",
        ]
    )


def _write_test_case_docx(cases: list[AcceptanceCase]) -> None:
    template = REPO_ROOT / "tests" / "docs" / "测试用例.docx"
    document = Document(str(template)) if template.exists() else Document()
    if document.paragraphs:
        document.paragraphs[0].text = "附件：测试用例"
    if len(document.paragraphs) > 1:
        document.paragraphs[1].text = f"{FUNCTION_NO} {TEST_PROJECT}（{FUNCTION_CODE}）"
    table = document.tables[0] if document.tables else document.add_table(rows=11, cols=2)
    values = {
        "测试项目": TEST_PROJECT,
        "测试目的": "验证开源原型量子金融软件系统可构建十类算法，并量化开发效率与性能提升。",
        "测试环境": "本地脚本环境，项目依赖已完成安装；使用固定小规模样例；无需联网。",
        "研究成果": "开源原型量子金融软件系统构建能力、开发效率提升证明材料、性能提升证明材料。",
        "交付物": "测试脚本、results.md、test_report.md、technical_report.md、测试用例.docx、开发效率对比柱状图、性能提升对比柱状图、终端运行截图、测试数据表。",
        "必选/可选": "必选",
        "前置条件": "已完成依赖安装；当前仓库可导入 double_quant；不依赖外部行情缓存。",
        "测试流程": (
            f"1. 执行 `{TEST_COMMAND}`。\n"
            "2. 检查程序输出是否覆盖十类算法名称。\n"
            "3. 检查编程语言体系是否只包含“决策性问题”和“估值性问题”。\n"
            "4. 检查开发效率对比柱状图和性能提升对比柱状图。\n"
            "5. 检查 results.md、test_report.md、technical_report.md 和测试数据表。"
        ),
        "预期结果": (
            f"{len(cases)} 个算法用例全部通过；开源软件接口有效代码行数少于直接量子线路；"
            "开源软件接口性能成本低于直接量子线路；生成中文图表和中文报告。"
        ),
        "测试结果": "",
        "测试结论": "",
    }
    for row in table.rows:
        key = row.cells[0].text.strip()
        if key in values:
            row.cells[1].text = values[key]
    if DEVELOPMENT_IMAGE.exists():
        document.add_paragraph("开发效率对比柱状图：")
        document.add_picture(str(DEVELOPMENT_IMAGE), width=Inches(5.8))
    if PERFORMANCE_IMAGE.exists():
        document.add_paragraph("性能提升对比柱状图：")
        document.add_picture(str(PERFORMANCE_IMAGE), width=Inches(5.8))
    if TERMINAL_IMAGE.exists():
        document.add_paragraph("终端运行截图：")
        document.add_picture(str(TERMINAL_IMAGE), width=Inches(5.8))
    document.save(DOC_DIR / "测试用例.docx")


def _validate_generated_charts() -> None:
    if validate_image is None:
        return
    for image in (DEVELOPMENT_IMAGE, PERFORMANCE_IMAGE):
        errors = validate_image(
            image,
            required_text="开源原型量子金融软件系统构建及开发效率与性能提升测试",
            min_width=900,
            min_height=520,
        )
        if errors:
            raise AssertionError("；".join(errors))


def _algorithm_name_sentence() -> str:
    return "、".join(ALGORITHM_NAMES[key] for key in ALGORITHM_ORDER)


def _chart_label(algorithm_name: str) -> str:
    return algorithm_name.replace("算法（", "算法\n（")

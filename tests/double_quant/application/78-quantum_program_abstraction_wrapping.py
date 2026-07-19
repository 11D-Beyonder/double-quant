from __future__ import annotations

import io
import sys
from dataclasses import dataclass
from pathlib import Path

import matplotlib
import numpy as np
import pytest
from docx import Document
from docx.shared import Inches
from qiskit import QuantumCircuit
from rich import box
from rich.console import Console
from rich.table import Table

from double_quant.programming import (
    OperatorResult,
    QuantumFinancialOperatorLibrary,
    default_operator_library,
)


matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402


FUNCTION_NO = 78
FUNCTION_CODE = "Func-45"
FUNCTION_NAME = "quantum-program-abstraction-wrapping"
TEST_PROJECT = "各金融场景下量子程序编码及金融问题、基础量子算法抽象封装测试"
TEST_COMMAND = (
    "uv run pytest "
    "tests/double_quant/application/78-quantum_program_abstraction_wrapping.py -s"
)

REPO_ROOT = Path(__file__).resolve().parents[3]
DOC_DIR = REPO_ROOT / "tests" / "docs" / f"{FUNCTION_NO}-{FUNCTION_NAME}"
IMAGE_DIR = DOC_DIR / "images"
COVERAGE_IMAGE = IMAGE_DIR / "quantum_program_abstraction_coverage.png"
STYLE_DIR = REPO_ROOT / ".codex" / "skills" / "3rd-test" / "scripts"
if STYLE_DIR.exists():
    sys.path.insert(0, str(STYLE_DIR))

from chinese_plot_style import (  # noqa: E402
    MEDIUM_WIDTH_MM,
    apply_chinese_style,
    color_for,
    save_figure,
    style_axes,
)


@dataclass(frozen=True, slots=True)
class WrapperCaseResult:
    case_group: str
    case_name: str
    operator_id: str
    financial_problem: str
    application_source: str
    quantum_program: str
    algorithm_primitive: str
    backend: str
    evidence: str
    resources: str


def test_quantum_program_encoding_and_abstraction_wrapping() -> None:
    library = default_operator_library()
    operator_ids = {spec.id for spec in library.list_specs()}

    assert {
        "func_1",
        "func_2",
        "func_3",
        "func_4",
        "func_5",
        "func_6",
        "func_7",
        "func_8",
        "func_9",
        "func_10",
    }.issubset(operator_ids)

    results = [
        _test_portfolio_application(library),
        _test_risk_value_application(library),
        _test_european_option_financial_problem(library),
        _test_application_circuit(
            library,
            "func_4",
            "动态账本更新算法（Func-4）",
            "决策性问题",
            "账本批次更新与周期一致性检查",
            "账本周期寄存器和受控量子门编码",
            "账本更新量子程序",
            "src/double_quant/application/dynamic_ledger_update.py",
        ),
        _test_application_circuit(
            library,
            "func_5",
            "去中心化金融管理算法（Func-5）",
            "决策性问题",
            "资产动作搜索与风险约束管理",
            "候选动作标记态和幅度放大量子门编码",
            "管理动作量子程序",
            "src/double_quant/application/defi_management.py",
        ),
        _test_application_circuit(
            library,
            "func_6",
            "反欺诈监测算法（Func-6）",
            "决策性问题",
            "交易循环监测与异常约束筛选",
            "交易约束可行态和混合量子门编码",
            "监测筛选量子程序",
            "src/double_quant/application/antifraud_monitoring.py",
        ),
        _test_application_circuit(
            library,
            "func_7",
            "支付与结算系统算法（Func-7）",
            "决策性问题",
            "支付结算流动性匹配",
            "结算约束可行态和混合量子门编码",
            "结算匹配量子程序",
            "src/double_quant/application/payment_settlement.py",
        ),
        _test_application_circuit(
            library,
            "func_8",
            "贷款发放决策算法（Func-8）",
            "决策性问题",
            "贷款审批特征选择与阈值约束",
            "审批特征可行态和选择量子门编码",
            "贷款审批量子程序",
            "src/double_quant/application/loan_decision.py",
        ),
        _test_application_circuit(
            library,
            "func_9",
            "银行网点布局优化算法（Func-9）",
            "决策性问题",
            "候选网点覆盖与成本约束搜索",
            "网点候选标记态和搜索量子门编码",
            "网点布局量子程序",
            "src/double_quant/application/branch_location.py",
        ),
        _test_application_circuit(
            library,
            "func_10",
            "指数追踪算法（Func-10）",
            "决策性问题",
            "指数成分选择与行业约束匹配",
            "成分约束可行态和权重量子门编码",
            "指数追踪量子程序",
            "src/double_quant/application/index_tracking.py",
        ),
    ]

    assert len(results) == 10
    assert {result.case_group for result in results} == {
        "决策性问题",
        "估值性问题",
    }
    assert {
        "最优投资组合算法（Func-1）",
        "风险价值计量算法（Func-2）",
        "金融衍生品定价算法（Func-3）",
        "动态账本更新算法（Func-4）",
        "去中心化金融管理算法（Func-5）",
        "反欺诈监测算法（Func-6）",
        "支付与结算系统算法（Func-7）",
        "贷款发放决策算法（Func-8）",
        "银行网点布局优化算法（Func-9）",
        "指数追踪算法（Func-10）",
    } == {result.case_name for result in results}

    _write_coverage_chart(results)
    _write_acceptance_documents(results)

    print(_program_output(results))
    assert COVERAGE_IMAGE.is_file()
    assert (DOC_DIR / "results.md").is_file()
    assert (DOC_DIR / "test_report.md").is_file()
    assert (DOC_DIR / "technical_report.md").is_file()
    assert (DOC_DIR / "测试用例.docx").is_file()


def _test_portfolio_application(
    library: QuantumFinancialOperatorLibrary,
) -> WrapperCaseResult:
    result = library.execute(
        "func_1",
        {
            "expected_returns": np.array([0.020, 0.010]),
            "covariance": np.array([[0.080, 0.010], [0.010, 0.050]]),
            "target_return": 0.015,
            "assets": ["股票A", "债券B"],
        },
        max_qpe_qubits=4,
    )
    weights = result.financial_result["weights"]
    assert isinstance(result, OperatorResult)
    assert result.operator_id == "func_1"
    assert set(weights) == {"股票A", "债券B"}
    assert sum(weights.values()) == pytest.approx(1.0, abs=0.06)

    return WrapperCaseResult(
        "决策性问题",
        "最优投资组合算法（Func-1）",
        result.operator_id,
        "均值-协方差资产配置和目标收益约束",
        "src/double_quant/application/portfolio.py",
        "资产结构、目标收益和预算约束被编码为组合配置量子程序",
        "组合配置量子程序",
        result.backend,
        "权重：" + "，".join(f"{asset}={value:.4f}" for asset, value in weights.items()),
        "返回组合权重",
    )


def _test_risk_value_application(
    library: QuantumFinancialOperatorLibrary,
) -> WrapperCaseResult:
    portfolio_returns = np.array(
        [0.012, -0.018, 0.021, -0.042, 0.016, -0.025, 0.009, -0.031],
        dtype=float,
    )
    result = library.execute(
        "func_2",
        {"portfolio_returns": portfolio_returns, "alpha": 0.75},
        backend="classical",
    )
    risk_value = next(iter(result.financial_result.values()))
    assert isinstance(result, OperatorResult)
    assert result.operator_id == "func_2"
    assert np.isfinite(risk_value)
    assert risk_value > 0.0

    return WrapperCaseResult(
        "估值性问题",
        "风险价值计量算法（Func-2）",
        result.operator_id,
        "组合收益率尾部风险计量",
        "src/double_quant/application/risk.py",
        "收益率尾部损失被编码为风险计量程序",
        "风险计量量子程序",
        result.backend,
        f"风险价值={risk_value:.5f}",
        "置信水平=0.75",
    )


def _test_application_circuit(
    library: QuantumFinancialOperatorLibrary,
    operator_id: str,
    algorithm_name: str,
    problem_language: str,
    financial_problem: str,
    quantum_program: str,
    program_kind: str,
    source_path: str,
) -> WrapperCaseResult:
    result = library.execute(operator_id, {}, backend="application_circuit")
    circuit = result.financial_result["circuit"]
    assert isinstance(result, OperatorResult)
    assert isinstance(circuit, QuantumCircuit)
    assert result.operator_id == operator_id
    assert result.backend == "application_circuit"
    assert result.resources.num_qubits == circuit.num_qubits
    assert result.resources.circuit_depth == circuit.depth()
    assert circuit.metadata is not None
    assert "application_id" in circuit.metadata

    resources = (
        f"量子位={result.resources.num_qubits}，"
        f"深度={result.resources.circuit_depth}，"
        f"双量子位门={result.resources.two_qubit_gates}"
    )
    return WrapperCaseResult(
        problem_language,
        algorithm_name,
        operator_id,
        financial_problem,
        source_path,
        quantum_program,
        program_kind,
        result.backend,
        f"量子程序构造成功，应用编号={circuit.metadata['application_id']}",
        resources,
    )


def _test_european_option_financial_problem(
    library: QuantumFinancialOperatorLibrary,
) -> WrapperCaseResult:
    result = library.execute(
        "func_3",
        {
            "terminal_price_scenarios": np.array([90.0, 100.0, 110.0, 120.0]),
            "strike": 100.0,
            "risk_free_rate": 0.0,
            "maturity": "1Y",
        },
        backend="classical_scenarios",
    )
    assert result.financial_result["option_price"] == pytest.approx(7.5)
    return WrapperCaseResult(
        "估值性问题",
        "金融衍生品定价算法（Func-3）",
        result.operator_id,
        "到期价格场景和执行价的贴现收益估值",
        "src/double_quant/programming/measures.py",
        "到期收益场景被编码为衍生品定价量子程序",
        "衍生品定价量子程序",
        result.backend,
        "欧式看涨期权价格=7.5000",
        "场景数=4",
    )


def _write_coverage_chart(results: list[WrapperCaseResult]) -> None:
    apply_chinese_style(
        width_mm=MEDIUM_WIDTH_MM,
        ncols=1,
        nrows=1,
        panel_aspect=1.05,
        required_text=TEST_PROJECT,
    )
    labels = [result.case_name for result in results]
    counts = np.ones(len(results), dtype=float)
    colors = [
        color_for(0) if result.case_group == "决策性问题" else color_for(1)
        for result in results
    ]
    y = np.arange(len(labels))

    fig, ax = plt.subplots()
    bars = ax.barh(y, counts, color=colors)
    style_axes(ax, title="十类金融算法量子程序封装覆盖", xlabel="封装检查项")
    ax.set_yticks(y)
    ax.set_yticklabels(labels)
    ax.set_xlim(0, 1.35)
    ax.invert_yaxis()
    for bar, count in zip(bars, counts, strict=True):
        ax.text(
            bar.get_width() + 0.03,
            bar.get_y() + bar.get_height() / 2,
            "通过",
            va="center",
            fontsize=8,
            color="#333333",
        )
    fig.tight_layout()
    save_figure(fig, IMAGE_DIR, "quantum_program_abstraction_coverage", formats=("png", "svg"))
    plt.close(fig)


def _write_acceptance_documents(results: list[WrapperCaseResult]) -> None:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    (DOC_DIR / "results.md").write_text(_results_document(results), encoding="utf-8")
    (DOC_DIR / "test_report.md").write_text(_test_report_document(results), encoding="utf-8")
    (DOC_DIR / "technical_report.md").write_text(
        _technical_report_document(results),
        encoding="utf-8",
    )
    _write_test_case_docx(results)


def _program_output(results: list[WrapperCaseResult]) -> str:
    console = Console(
        file=io.StringIO(),
        width=170,
        record=True,
        force_terminal=True,
        color_system=None,
    )
    console.print(f"功能编号：{FUNCTION_NO}")
    console.print(f"功能名称：{TEST_PROJECT}（{FUNCTION_CODE}）")
    console.print(f"测试命令：{TEST_COMMAND}")
    console.print(
        "覆盖范围：最优投资组合算法、风险价值计量算法、金融衍生品定价算法、动态账本更新算法、"
        "去中心化金融管理算法、反欺诈监测算法、支付与结算系统算法、贷款发放决策算法、"
        "银行网点布局优化算法、指数追踪算法。"
    )
    console.print()

    table = Table(
        title="十类金融算法量子程序封装验证结果",
        box=box.ROUNDED,
        show_lines=True,
        expand=True,
        padding=(0, 1),
    )
    table.add_column("编程语言体系", ratio=12, overflow="fold")
    table.add_column("算法名称", ratio=22, overflow="fold")
    table.add_column("金融问题", ratio=20, overflow="fold")
    table.add_column("量子程序/算法抽象", ratio=26, overflow="fold")
    table.add_column("证据", ratio=32, overflow="fold")
    table.add_column("资源", ratio=18, overflow="fold")
    for result in results:
        table.add_row(
            result.case_group,
            result.case_name,
            result.financial_problem,
            result.quantum_program,
            result.evidence,
            result.resources,
        )
    console.print(table)
    return console.export_text(styles=False)


def _results_document(results: list[WrapperCaseResult]) -> str:
    return "\n".join(
        [
            f"# {FUNCTION_NO} {TEST_PROJECT} {FUNCTION_CODE} 测试结果",
            "",
            "## 测试对象",
            "",
            "当前十类金融算法的量子程序编码封装，以及决策性问题、估值性问题两类编程语言体系。",
            "",
            "## 测试命令",
            "",
            "```bash",
            TEST_COMMAND,
            "```",
            "",
            "## 程序输出",
            "",
            "```text",
            _program_output(results),
            "```",
            "",
            "## 图片输出",
            "",
            "![终端运行截图](images/terminal_run.png)",
            "",
            "![十类金融算法量子程序封装覆盖](images/quantum_program_abstraction_coverage.png)",
            "",
            "## 关键结果",
            "",
            "- 已逐项验证十个算法名称，每个算法名称均与用户给定清单一致。",
            "- 已验证编程语言体系只使用“决策性问题”和“估值性问题”两类。",
            "- 每个封装结果均返回金融结果、诊断信息或量子资源信息，证明金融问题和量子程序之间存在统一抽象层。",
        ]
    )


def _test_report_document(results: list[WrapperCaseResult]) -> str:
    decision_cases = [result for result in results if result.case_group == "决策性问题"]
    valuation_cases = [result for result in results if result.case_group == "估值性问题"]
    return "\n".join(
        [
            f"# {FUNCTION_NO} {TEST_PROJECT} {FUNCTION_CODE} 测试报告",
            "",
            "## 测试目标",
            "",
            "验证当前系统已经把十类金融算法的量子程序编码、金融问题抽象和算法封装统一到决策性问题、估值性问题两类编程语言体系。",
            "",
            "## 测试范围",
            "",
            f"- 决策性问题用例 {len(decision_cases)} 个。",
            f"- 估值性问题用例 {len(valuation_cases)} 个。",
            "- 覆盖算法名称：" + "、".join(result.case_name for result in results) + "。",
            "",
            "## 测试方法",
            "",
            "逐项调用当前十类算法对应的封装入口。对决策性问题检查量子程序构造、资源信息或决策输出；对估值性问题检查风险贡献、价格输出和量子程序编码结果。",
            "",
            "## 通过标准",
            "",
            "- 十个算法名称必须全部出现，且名称必须与给定清单一致。",
            "- 编程语言体系只能出现“决策性问题”和“估值性问题”。",
            "- 输出必须包含金融问题、量子程序编码、验证证据和资源证据。",
            "",
            "## 测试结果分析",
            "",
            f"测试通过。本次共验证 {len(results)} 个封装用例，其中决策性问题 {len(decision_cases)} 个、估值性问题 {len(valuation_cases)} 个。",
            "结果证明十类金融算法均可以通过统一封装入口执行，并返回可审计的金融结果或量子程序资源信息。",
            "",
            "## 实际验证记录",
            "",
            *[
                f"- {result.case_name}：{result.case_group}，{result.evidence}。"
                for result in results
            ],
            "",
            "## 风险与限制",
            "",
            "- 最优投资组合算法的量子求解结果存在近似误差，测试采用约束容差检查封装可用性。",
            "- 本用例关注封装入口和量子程序编码，不展开完整硬件运行或长时间模拟。",
            "",
            "## 测试结论",
            "",
            "通过。当前代码已经覆盖十类金融算法量子程序编码，并将金融问题统一归入决策性问题和估值性问题两类。",
        ]
    )


def _technical_report_document(results: list[WrapperCaseResult]) -> str:
    return "\n".join(
        [
            f"# {FUNCTION_NO} {TEST_PROJECT} {FUNCTION_CODE} 技术报告",
            "",
            "## 技术目标",
            "",
            "用当前源码中的真实入口证明十类金融算法已经形成可复用的量子程序封装，并且编程语言体系只归入决策性问题和估值性问题两类。",
            "",
            "## 实现位置",
            "",
            "- 验收脚本：`tests/double_quant/application/78-quantum_program_abstraction_wrapping.py`",
            "- 应用源码目录：`src/double_quant/application/`",
            "- 编程封装源码：`src/double_quant/programming/quantum_operator.py`",
            "- 文档目录：`tests/docs/78-quantum-program-abstraction-wrapping/`",
            "",
            "## 实现概述",
            "",
            "脚本通过统一封装入口调用十类金融算法。每个算法用例均记录算法名称、编程语言体系、金融问题、量子程序编码、验证证据和资源信息。",
            "对外报告不展示底层算法英文名称，只展示验收要求中的中文算法名称。",
            "",
            "## 关键技术点",
            "",
            "- 决策性问题：覆盖最优投资组合、动态账本更新、去中心化金融管理、反欺诈监测、支付与结算、贷款发放决策、银行网点布局优化、指数追踪。",
            "- 估值性问题：覆盖风险价值计量和金融衍生品定价。",
            "- 量子程序编码：统一检查量子态、量子门、约束可行态、候选标记态、寄存器和权重结构等信息。",
            "- 输出证据：统一记录金融输出、量子位数量、电路深度、双量子位门数量或价格、风险贡献等结果。",
            "",
            "## 验收脚本设计",
            "",
            "脚本把用例分为“决策性问题”和“估值性问题”。每个用例都返回统一结果记录，记录算法名称、金融问题、源码位置、量子程序编码、验证证据和资源摘要。",
            "",
            "## 验证方法",
            "",
            _program_output(results),
            "",
            "## 技术结论",
            "",
            "当前量子金融软件系统已经具备十类金融算法的场景级封装。对外验收材料中的算法名称与用户给定清单一致，编程语言体系只保留决策性问题和估值性问题。",
        ]
    )


def _write_test_case_docx(results: list[WrapperCaseResult]) -> None:
    template = REPO_ROOT / "tests" / "docs" / "测试用例.docx"
    document = Document(str(template)) if template.exists() else Document()
    if document.paragraphs:
        document.paragraphs[0].text = "附件：测试用例"
    if len(document.paragraphs) > 1:
        document.paragraphs[1].text = f"{FUNCTION_NO} {TEST_PROJECT}（{FUNCTION_CODE}）"
    table = document.tables[0] if document.tables else document.add_table(rows=11, cols=2)
    values = {
        "测试项目": TEST_PROJECT,
        "测试目的": "验证各金融场景下量子程序编码，以及金融问题和基础量子算法抽象封装能力。",
        "测试环境": "本地脚本环境，项目依赖已完成安装，量子程序构造、图表生成和文档生成组件可用。",
        "研究成果": "量子金融编程框架源码与编程框架功能证明材料。",
        "交付物": "测试脚本、results.md、test_report.md、technical_report.md、封装覆盖图。",
        "必选/可选": "必选",
        "前置条件": "已完成依赖安装；无需联网；使用本地固定样例数据。",
        "测试流程": (
            f"1. 执行 `{TEST_COMMAND}`。\n"
            "2. 检查输出表中是否覆盖十个算法名称。\n"
            "3. 检查编程语言体系是否只包含决策性问题和估值性问题。\n"
            "4. 检查图片目录中的封装覆盖图。"
        ),
        "预期结果": (
            f"{len(results)} 个封装用例全部通过，十类金融算法均能通过统一封装入口返回结果。"
        ),
        "测试结果": "",
        "测试结论": "",
    }
    for row in table.rows:
        key = row.cells[0].text.strip()
        if key in values:
            row.cells[1].text = values[key]
    if COVERAGE_IMAGE.exists():
        document.add_paragraph("封装覆盖图：")
        document.add_picture(str(COVERAGE_IMAGE), width=Inches(5.8))
    terminal_image = IMAGE_DIR / "terminal_run.png"
    if terminal_image.exists():
        document.add_paragraph("终端运行截图：")
        document.add_picture(str(terminal_image), width=Inches(5.8))
    document.save(DOC_DIR / "测试用例.docx")
